"""
가이드 .docx 파일을 파싱해서 텍스트로 추출하는 모듈
"""
from docx import Document
import io


def extract_guide_text(file_bytes: bytes) -> str:
    """
    업로드된 .docx 파일 바이트를 받아서 표/문단 텍스트를 전부 추출.
    표 형식 가이드(흑돈바라기 가이드 같은 구조)도 안전하게 처리.
    """
    doc = Document(io.BytesIO(file_bytes))
    lines = []

    # 1) 일반 문단
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            lines.append(text)

    # 2) 표 안의 내용 (가이드가 보통 표 형식이라 이 부분이 핵심)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            # 빈 셀 제거하고 " | "로 합치기
            cells = [c for c in cells if c]
            if cells:
                lines.append(" | ".join(cells))

    return "\n".join(lines)


if __name__ == "__main__":
    import sys
    with open(sys.argv[1], "rb") as f:
        content = f.read()
    print(extract_guide_text(content))
